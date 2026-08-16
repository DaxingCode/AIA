# -*- coding: utf-8 -*-
"""
商业计划书生成脚本
产品：好记AI —— 一人主导的AI个人生活助理（双端融合）
赛事：2026年南宁市人工智能OPC创业大赛
输出：DOCX + PDF（同一份结构化内容，两种渲染）
用法：python generate_bp.py
"""
import os

# ===================== 调色板 =====================
PRIMARY = "1D4ED8"   # 主色 蓝
ACCENT  = "0EA5E9"   # 亮蓝
TEAL    = "14B8A6"   # 青绿（健康/AI）
DARK    = "1F2937"   # 正文深灰
GRAY    = "6B7280"   # 次要灰
LIGHT   = "F1F5F9"   # 浅底
FONT_CN = "Microsoft YaHei"

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "AI助理商业计划书_2026南宁OPC大赛.docx")
PDF_PATH  = os.path.join(OUT_DIR, "AI助理商业计划书_2026南宁OPC大赛.pdf")

COVER = {
    "title": "AI助理·好记AI",
    "subtitle": "自动记账记待办，管理饮食和健康（小程序、App双端融合）",
    "event": "2026年南宁市人工智能OPC创业大赛 · 参赛项目商业计划书",
    "applicant": "参赛人 / OPC 核心人物：韦文幸",
    "category": "参赛类别：AI应用达人 / AI集成达人",
    "date": "2026年8月",
}

# 章节目录（用于生成目录与页码）
TOC = [
    "执行摘要",
    "一、项目背景与痛点",
    "二、产品与服务（双端融合）",
    "三、核心技术与创新点",
    "四、市场分析",
    "五、商业模式",
    "六、营销策略",
    "七、团队与 OPC 创业特质",
    "八、落地可行性与发展规划",
    "九、财务预测（测算）",
    "十、风险分析与对策",
    "十一、社会价值与政策契合",
    "十二、参赛承诺与落地计划",
    "附录：知识产权与技术支撑",
]

# ===================== 内容模型 =====================
# 元素类型：h1 / h2 / p / bullets / table / callout / pagebreak
DOC = []

def h1(t): DOC.append({"t": "h1", "v": t})
def h2(t): DOC.append({"t": "h2", "v": t})
def p(t):  DOC.append({"t": "p", "v": t})
def bullets(items): DOC.append({"t": "bullets", "v": items})
def table(headers, rows, widths=None): DOC.append({"t": "table", "v": {"h": headers, "r": rows, "w": widths}})
def callout(t): DOC.append({"t": "callout", "v": t})
def pagebreak(): DOC.append({"t": "pagebreak", "v": ""})

# ===================== 正文内容 =====================
h1("执行摘要")
p("好记AI 是我一个人做的（OPC 模式）AI 生活小助手。它有两个入口：iPhone 上的「好记AI」App 和微信里的「好好吃饭」小程序，两边数据打通，一起管你的记账、吃饭、健康和待办这四件事。你只要截图、拍照，或者把健康数据接进来，它就自动帮你记好；吃进多少、消耗多少，也能对上账。")
callout("四个最实在的亮点：① 一个助手管四件事，手机 App 和小程序两边都能用；② 截图、拍照就能记，不用手动打字；③ 吃的和动的能连起来算，知道每天到底是攒了还是亏了；④ 我一个人就能做出来，靠 AI 工具帮忙，成本低、能复制，正好贴合这次 OPC 一人公司的主题。")
p("为什么现在做这件事：现在国家大力推“人工智能+”，大家都想用 AI 把自己的生活管好。但市面上的 App 太碎了——记账的只管钱，管饭的只管吃，健康的和待办的又是各管各的，谁也不搭谁。好记AI 就是把这四件事用 AI 串起来，正好踩在“AI 进到生活里”这个点上。")
p("关于我怎么做的：整个项目就我一个人（韦文幸）在主导。我用 AI 写代码工具、云端函数代理，再按需找设计外包，把 iOS App、小程序和云端后台都做了出来。启动成本很低、速度快，做法也能复制，这就是典型的 OPC 一人公司。从截图识别到热量对账的功能，我已经在真机上全部跑通验证过了，项目是能落地的，我也希望能落地南宁、加入 OPC 创业社区。")

h1("一、项目背景与痛点")
h2("1.1 时代背景")
p("到了 2026 年，“人工智能+”已经不只是口号，AI 正在变成像水电一样的基础能力。普通人最实在的一个念头就是：能不能让 AI 帮我把生活打理好——钱花哪了、吃了些啥、身体指标怎么样、还有啥事没办。可这些事，一直散在各个 App 里，没人把它们放一块儿。")
h2("1.2 用户痛点")
bullets([
    "东西太分散：记账、吃饭、健康、待办分别在好几个 App 里，每天来回切，根本拼不出一张完整的“吃了多少、消耗多少”的图。",
    "记起来太麻烦：以前都得自己手打金额、食物重量、待办事项，门槛高，没几天就放弃了。",
    "怕隐私泄露：钱和身体数据都很私密，大家一般不愿意把原图、原数据直接传到网上。",
    "算得不准：光靠大模型猜食物热量，误差能到正负三成，没有标准营养库兜底，根本没法拿来管健康。",
])
h2("1.3 市场已被验证")
p("这个方向已经被验证过了：做截图记账的咔皮记账（商汤做的），2026 年 1 月用户就过了 500 万，说明“截图就能记账”是真需求；健康类 App 用户更是上亿。需求是实打实存在的，缺的只是把“记账、吃饭、健康、待办”合在一起、还能形成闭环的那个人。")

h1("二、产品与服务（双端融合）")
p("好记AI 是“一个产品、两个入口”：iPhone App 负责系统级的感知和全场景助理，微信小程序专门做饮食这一块，借着微信好传播、获客便宜。两个端通过云端实时同步，数据始终一致。")
h2("2.1 iOS App「好记AI」（原生 Swift + SwiftUI）")
bullets([
    "截图自动记：iPhone 一截屏，快捷指令就无感触发，云端 AI 把图变成结构化的账单、待办或知识卡，再用本地通知汇总给你。",
    "拍照算食物热量：对着饭拍一张，AI 加上标准营养库校准，直接给你名称、重量、热量，还有蛋白质、碳水、脂肪。",
    "HealthKit 健康数据：把运动、睡眠、身体指标拉进来，做成健康看板和趋势图。",
    "热量能对账：用运动和身体指标算出每天总消耗，减去吃进去的热量，得出净热量，吃和动就连起来了。",
    "云同步：换设备数据也不丢，两端始终对得上。",
])
h2("2.2 微信小程序「好好吃饭」（饮食垂类双端承载）")
bullets([
    "文字或拍照录饮食，recognizeFood 云函数识别，复用共享的食物库。",
    "7 天饮食查询页，带营养小格子（碳水、蛋白质、脂肪、糖、膳食纤维、钠）。",
    "包装食品没标重量时，默认按 100 克算，你也能改重量重新算。",
    "能“分享到朋友圈或好友”，靠社交传播便宜地拉新。",
])
h2("2.3 双端融合架构")
table(
    ["层级", "iOS App「好记AI」", "微信小程序「好好吃饭」", "统一后端"],
    [
        ["采集入口", "快捷指令 / 相机 / HealthKit", "文字 / 拍照 / 分享", "—"],
        ["AI 识别", "多模态 LLM → JSON", "recognizeFood 云函数", "CloudBase 云函数代理"],
        ["数据层", "本地优先 SwiftData", "云数据库 foodDB", "共享营养库"],
        ["同步", "aia-sync 双向同步", "aia-sync 双向同步", "统一 sync 路由"],
        ["呈现", "净热量闭环 / 看板", "7日饮食 / 营养网格", "—"],
    ],
    widths=[0.16, 0.30, 0.30, 0.24],
)
callout("我们的壁垒：单独做某一块的产品不少，但把“截图记账 + 食物热量 + 健康数据 + 闭环对账”四样合在一起的，目前还几乎没有，这正是好记AI 能切入的地方。")

h1("三、核心技术与创新点")
bullets([
    "多模态识别引擎：图片传到云端大模型，再变成结构化的 JSON；我们用 prompt 把返回字段卡死，下游能直接存库。",
    "营养库校准：大模型估热量能有正负三成误差，我们接了标准营养库（USDA、薄荷健康、Keep）做“菜名→标准热量”的映射，准多了。",
    "本地优先的隐私设计：钱和健康数据默认存在手机本地；要上云也只传整理好的字段，不传原图；密钥走 CloudBase 云函数代理，前端不写死。",
    "热量闭环算法：吃进去的和消耗掉的，在同一个人的数据库里联动，这是和所有单点产品最大的不同。",
    "一个人就能开发：靠 AI 写代码工具 + 云函数代理 + 按需设计外包，一个人把 iOS、小程序和云端全做了；这套流程能复制、能放大。",
])

h1("四、市场分析")
h2("4.1 目标用户")
bullets([
    "上班族和学生：图省事，想截图就记账、拍照就记录，少动手。",
    "减脂、健身的人：在意热量和营养搭配，需要“吃和动”能对上。",
    "关注慢病的人：要长期盯住饮食和身体指标。",
    "家庭和老人：需要门槛特别低的方式（截图、拍照代替打字）。",
])
h2("4.2 竞品对比")
table(
    ["品类", "代表产品", "主要短板", "本项目差异"],
    [
        ["截图记账", "咔皮记账 / 蜜蜂记账", "仅账单，无健康/待办", "四合一 + 净热量闭环"],
        ["截图整理", "RecallBox / CaptainShot", "国内支付弱、无健康", "本土支付 + 健康联动"],
        ["纯 OCR", "全能图片转文字", "仅取字，无结构化", "识别即结构化落库"],
        ["饮食", "薄荷健康 / Keep", "无截图记账/待办闭环", "双端 + 闭环"],
        ["健康", "各类运动健康 App", "无饮食摄入联动", "摄入—消耗联动"],
    ],
    widths=[0.16, 0.26, 0.30, 0.28],
)
h2("4.3 市场规模（测算）")
p("健康管理类 App 国内用户上亿，AI 应用市场还在快速增长；截图记账头部产品用户已经破 500 万。把“想省事”和“想管好健康”这两拨人叠起来看，潜在用户是千万级别的，而且大家越重视健康，越愿意花钱。上面是按公开标杆估的，不是精确统计。")

h1("五、商业模式")
h2("5.1 收入来源")
bullets([
    "会员订阅 Pro（约 ¥18/月 或 ¥128/年）：高级识别额度、无限云同步、深度健康报告、去广告。",
    "增值识别包：超出免费额度的 AI 识别次数按量计费。",
    "广告位（已有 AdBanner 云控模块）：信息流 / 激励广告，按展示与点击结算。",
    "B 端能力输出：识别 API 与营养数据服务向中小开发者开放。",
    "数据洞察报告：周期性健康 / 消费洞察（脱敏聚合）增值服务。",
])
h2("5.2 成本结构（轻资产）")
p("我们几乎没有固定成本：没有自己的服务器，云端函数和 AI 接口按量付费，存储用托管数据库，研发就我一个人加按需外包。用得越多成本才涨一点，规模上来反而更划算。")
h2("5.3 单位经济与 OPC 优势")
p("一个用户一个月给我们带来的收入 = 会员费 + 广告 + 增值识别；成本 = 云函数 + API 调用。因为固定成本极低，哪怕用户量还不算大，一个人也能赚到钱——这正是 OPC 模式的好处：轻、回本快、能复制。")

h1("六、营销策略")
bullets([
    "落地南宁：加入 OPC 创业社区，去高校宣讲、进社区做健康记录活动，借比赛的流量把名气打出去。",
    "做好应用商店的展示和真实用户评价。",
    "内容引流：在小红书、抖音上讲“AI 怎么帮我省时间、管健康”，慢慢攒起我自己的个人号。",
    "小程序裂变：已经能做“分享到朋友圈/好友”，靠社交便宜拉新。",
    "社群和达人：跟减脂、健身的社群合作，用“拍张照就算出卡路里”的惊喜感带出口碑。",
])

h1("七、团队与 OPC 创业特质")
p("这项目就是典型的 One Person Company（OPC）——我一个人（韦文幸）主导。我本来就有微信小程序云开发的实战经验，正在系统学 Swift 原生开发，iOS 的 MVP 全部功能和小程序的饮食模块，都是我一个人在真机上跑通验证的。")
bullets([
    "靠 AI 帮忙：用 AI 写代码工具生成代码，用云函数代理搞定密钥和后台，视觉不够就按需找设计外包，一个人也能交付全栈。",
    "轻资产：没有固定办公室，也不必养一队人，启动就能跑。",
    "能复制：从需求 → AI 识别 → 结构化 → 双端呈现，这套流程已经成了模板，能很快挪到别的场景（发票、名片等）。",
    "外包协作：云端函数代理、设计外包都是按需调用，正好是“一个人主导 + 外包支撑”的 OPC 样子。",
])

h1("八、落地可行性与发展规划")
h2("8.1 已验证里程碑（真机验证通过）")
table(
    ["阶段", "内容", "状态"],
    [
        ["M1", "分享扩展导入截图 → 云端 LLM → 账单/待办/知识卡", "已完成"],
        ["M2", "快捷指令“截屏”自动化对接（无感触发）", "已完成"],
        ["M3", "拍照食物识别卡路里 + 营养库校正", "已完成"],
        ["M4", "HealthKit 运动/睡眠/身体指标看板", "已完成"],
        ["M5", "净热量闭环 + 云同步双向通", "已完成（真机验证）"],
    ],
    widths=[0.12, 0.66, 0.22],
)
h2("8.2 产品现状")
p("iPhone 上的「好记AI」App，MVP 的全部功能都已经在真机上验证过；微信里的「好好吃饭」小程序，饮食模块已经上线在跑了，两端通过 aia-sync 云函数实时同步。")
h2("8.3 发展路线图")
bullets([
    "近期：广告位精细化运营、会员体系上线、分享裂变放大。",
    "中期：上架后平滑切换 iCloud 同步；网页 / 安卓端复用 CloudBase 后端；扩展识别场景（票据、名片）。",
    "远期：开放平台 / 识别 API 输出；健康洞察 AI 顾问。",
])
h2("8.4 落地南宁计划")
p("我们想落地南宁，加入人工智能 OPC 创业社区，享受场地、水电、物业补贴，还有算力支持、创业担保贷款和投融资对接这些政策，也参加大赛的孵化和资源对接。")

h1("九、财务预测（测算）")
p("下面是基于“轻资产”这个前提做的测算示意，不是审计数据，只是用来说明赚钱的逻辑和增长走势。")
table(
    ["指标（测算）", "第 1 年", "第 2 年", "第 3 年"],
    [
        ["活跃用户（万人）", "3", "15", "50"],
        ["付费率", "4%", "6%", "8%"],
        ["年营业收入（万元）", "12", "90", "380"],
        ["年运营成本（万元）", "8", "40", "120"],
        ["年利润（万元）", "4", "50", "260"],
    ],
    widths=[0.34, 0.22, 0.22, 0.22],
)
p("几个关键假设：用户主要靠内容营销加小程序裂变、用很低的成本拉来；收入以会员订阅为主，广告和增值为辅；成本主要是云函数和 API 调用。照这样，第一年攒够付费用户就能打平。")

h1("十、风险分析与对策")
table(
    ["风险", "说明", "对策"],
    [
        ["技术适配", "iOS 系统版本/权限差异影响自动化体验", "引导页 + 手动兜底入口，覆盖全版本"],
        ["隐私合规", "财务/健康数据敏感", "本地优先 + 隐私政策 + HealthKit 合规声明"],
        ["识别精度", "LLM 估算热量误差", "营养库校正 + 用户确认页后再入库"],
        ["API 成本波动", "多模态调用按量计费", "云函数代理 + 缓存 + 额度控制"],
        ["竞争", "大厂入场复制单点功能", "闭环壁垒 + 双端融合 + 快速迭代"],
    ],
    widths=[0.18, 0.40, 0.42],
)

h1("十一、社会价值与政策契合")
bullets([
    "呼应国家“人工智能+”行动，是 AI 真正用进个人生活里的一个实在例子。",
    "给大家一个能复制的 OPC 创业样板，借着“以赛促创、以赛引投”，带动个人创业就业。",
    "服务“健康中国”：靠饮食和运动的闭环，帮着防慢病、管好健康。",
    "帮南宁的数字经济：这种轻资产、能孵化的模式，能催生出更多 OPC 项目。",
])

h1("十二、参赛承诺与落地计划")
p("我们同意大赛组委会对参赛项目做宣传展示、推广对接、孵化培育等活动，也承诺项目真实合规、权属清楚、没有知识产权纠纷。我们有意在南宁发展、加入 OPC 创业社区，积极配合资源对接和专家陪跑，争取落邕和入孵奖励，把项目真正做下去、做长久。")

h1("附录：知识产权与技术支撑")
h2("A. 知识产权布局")
bullets([
    "商标：已经做了 App 名称和 Logo 的原创性自查，也备了改名方案的检索清单。",
    "软著：打算给核心的识别和同步逻辑申请软件著作权。",
    "技术壁垒：本地优先的隐私设计、热量闭环算法、双端同步协议。",
])
h2("B. 技术栈清单")
bullets([
    "iOS：Swift + SwiftUI + SwiftData，HealthKit / PhotoKit / Camera 原生桥接。",
    "小程序：微信云开发（云函数 + 云数据库 + 云存储）。",
    "AI：云端多模态 LLM（结构化 JSON 输出）+ 标准营养数据库校正。",
    "后端：CloudBase 云函数 aia-sync / recognizeFood，API Key 经云函数代理。",
])
h2("C. 联系方式")
p("参赛人：韦文幸 ｜ 项目：好记AI ｜ 类别：AI应用达人 / AI集成达人")

# ===================== DOCX 渲染 =====================
def render_docx(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_run(run, size=None, bold=None, color=None, font=FONT_CN):
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        if size: run.font.size = Pt(size)
        if bold is not None: run.font.bold = bold
        if color: run.font.color.rgb = RGBColor.from_string(color)

    doc = Document()
    # 默认正文样式
    normal = doc.styles['Normal']
    normal.font.name = FONT_CN
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    # 页边距
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(0.95); s.right_margin = Inches(0.95)

    # 封面
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(COVER["title"]); set_run(r, 30, True, PRIMARY)
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(COVER["subtitle"]); set_run(r, 15, False, DARK)
    for _ in range(1): doc.add_paragraph()
    ev = doc.add_paragraph(); ev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ev.add_run(COVER["event"]); set_run(r, 12, True, ACCENT)
    for _ in range(6): doc.add_paragraph()
    for line in [COVER["applicant"], COVER["category"], COVER["date"]]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(line); set_run(r, 11, False, GRAY)
    doc.add_page_break()

    # 目录
    tc = doc.add_paragraph(); r = tc.add_run("目  录"); set_run(r, 16, True, PRIMARY)
    tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, item in enumerate(TOC, 1):
        pp = doc.add_paragraph(); pp.paragraph_format.space_after = Pt(4)
        r = pp.add_run(f"{i:>2}.  {item}"); set_run(r, 11, False, DARK)
    doc.add_page_break()

    # 正文
    for el in DOC:
        if el["t"] == "h1":
            para = doc.add_heading(level=1)
            r = para.add_run(el["v"]); set_run(r, 15, True, PRIMARY)
            para.paragraph_format.space_before = Pt(10); para.paragraph_format.space_after = Pt(6)
        elif el["t"] == "h2":
            para = doc.add_heading(level=2)
            r = para.add_run(el["v"]); set_run(r, 12.5, True, DARK)
            para.paragraph_format.space_before = Pt(6); para.paragraph_format.space_after = Pt(3)
        elif el["t"] == "p":
            para = doc.add_paragraph(); para.paragraph_format.space_after = Pt(6)
            r = para.add_run(el["v"]); set_run(r, 10.5, False, DARK); r.font.name = FONT_CN
        elif el["t"] == "bullets":
            for it in el["v"]:
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.space_after = Pt(3)
                r = para.add_run(it); set_run(r, 10.5, False, DARK)
        elif el["t"] == "callout":
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), LIGHT)
            pPr.append(shd)
            para.paragraph_format.left_indent = Inches(0.1)
            para.paragraph_format.space_before = Pt(4); para.paragraph_format.space_after = Pt(6)
            r = para.add_run("◆ " + el["v"]); set_run(r, 10.5, True, PRIMARY)
        elif el["t"] == "table":
            _render_docx_table(doc, el["v"], set_run)
        elif el["t"] == "pagebreak":
            doc.add_page_break()

    # 页脚页码
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("好记AI  |  2026年南宁市人工智能OPC创业大赛  |  第 ")
    set_run(run, 8, False, GRAY)
    # 简单页码域
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    fp._p.append(fld1)
    run = fp.add_run(" 页"); set_run(run, 8, False, GRAY)

    doc.save(path)


def _render_docx_table(doc, tv, set_run):
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    headers = tv["h"]; rows = tv["r"]; widths = tv.get("w")
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    # 表头
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]; r = para.add_run(htext)
        set_run(r, 10, True, "FFFFFF")
        # 表头底色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), PRIMARY)
        tcPr.append(shd)
    # 数据行
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            para = cells[ci].paragraphs[0]; r = para.add_run(val)
            set_run(r, 9.5, False, DARK)
            if ri % 2 == 1:
                tcPr = cells[ci]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), LIGHT)
                tcPr.append(shd)
    # 列宽
    if widths:
        for ci, w in enumerate(widths):
            for cell in tbl.columns[ci].cells:
                cell.width = Inches(w * 6.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ===================== PDF 渲染 =====================
def render_pdf(path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, PageBreak, HRFlowable)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    F = 'STSong-Light'
    C = lambda h: colors.HexColor('#' + h)

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle('title', fontName=F, fontSize=26, leading=32,
                             textColor=C(PRIMARY), alignment=TA_CENTER, spaceAfter=6)
    subt_s = ParagraphStyle('subt', fontName=F, fontSize=14, leading=20,
                            textColor=C(DARK), alignment=TA_CENTER)
    evt_s = ParagraphStyle('evt', fontName=F, fontSize=12, leading=18,
                           textColor=C(ACCENT), alignment=TA_CENTER, spaceBefore=10)
    meta_s = ParagraphStyle('meta', fontName=F, fontSize=11, leading=18,
                            textColor=C(GRAY), alignment=TA_CENTER)
    h1_s = ParagraphStyle('h1', fontName=F, fontSize=15, leading=21,
                          textColor=C(PRIMARY), spaceBefore=12, spaceAfter=6)
    h2_s = ParagraphStyle('h2', fontName=F, fontSize=12.5, leading=18,
                          textColor=C(DARK), spaceBefore=7, spaceAfter=3)
    p_s = ParagraphStyle('p', fontName=F, fontSize=10.5, leading=16.5,
                        textColor=C(DARK), spaceAfter=6, alignment=TA_LEFT)
    bullet_s = ParagraphStyle('bullet', fontName=F, fontSize=10.5, leading=16,
                              textColor=C(DARK), spaceAfter=3, leftIndent=12, bulletIndent=2)
    callout_s = ParagraphStyle('callout', fontName=F, fontSize=10.5, leading=16,
                               textColor=C(PRIMARY), spaceBefore=4, spaceAfter=6, leftIndent=8)
    toc_s = ParagraphStyle('toc', fontName=F, fontSize=11, leading=20, textColor=C(DARK))
    cell_h_s = ParagraphStyle('cellh', fontName=F, fontSize=10, leading=13, textColor=colors.white)
    cell_s = ParagraphStyle('cell', fontName=F, fontSize=9.5, leading=13, textColor=C(DARK))

    story = []

    # 封面
    story += [Spacer(1, 3.2*cm), Paragraph(COVER["title"], title_s),
              Paragraph(COVER["subtitle"], subt_s), Spacer(1, 0.6*cm),
              Paragraph(COVER["event"], evt_s), Spacer(1, 4.5*cm),
              Paragraph(COVER["applicant"], meta_s),
              Paragraph(COVER["category"], meta_s),
              Paragraph(COVER["date"], meta_s)]
    story.append(PageBreak())

    # 目录
    story.append(Paragraph("目  录", ParagraphStyle('toct', parent=title_s, fontSize=16, spaceAfter=12)))
    for i, item in enumerate(TOC, 1):
        story.append(Paragraph(f"{i:>2}.  {item}", toc_s))
    story.append(PageBreak())

    # 正文
    for el in DOC:
        if el["t"] == "h1":
            story.append(Paragraph(el["v"], h1_s))
        elif el["t"] == "h2":
            story.append(Paragraph(el["v"], h2_s))
        elif el["t"] == "p":
            story.append(Paragraph(el["v"].replace("&", "&amp;"), p_s))
        elif el["t"] == "bullets":
            for it in el["v"]:
                story.append(Paragraph("• " + it.replace("&", "&amp;"), bullet_s))
        elif el["t"] == "callout":
            story.append(HRFlowable(width="100%", thickness=0.6, color=C(TEAL), spaceBefore=2, spaceAfter=3))
            story.append(Paragraph("◆ " + el["v"].replace("&", "&amp;"), callout_s))
        elif el["t"] == "table":
            story.append(_render_pdf_table(el["v"], cell_h_s, cell_s, C))
            story.append(Spacer(1, 0.2*cm))
        elif el["t"] == "pagebreak":
            story.append(PageBreak())

    def footer(canvas, docobj):
        canvas.saveState()
        canvas.setFont(F, 8); canvas.setFillColor(C(GRAY))
        canvas.drawCentredString(A4[0]/2, 1.1*cm,
            "好记AI  |  2026年南宁市人工智能OPC创业大赛  |  第 %d 页" % docobj.page)
        canvas.restoreState()

    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=1.6*cm, bottomMargin=1.6*cm,
                            leftMargin=1.9*cm, rightMargin=1.9*cm,
                            title="AI助理商业计划书_2026南宁OPC大赛",
                            author="韦文幸")
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def _render_pdf_table(tv, cell_h_s, cell_s, C):
    from reportlab.lib.units import cm
    from reportlab.platypus import Table, TableStyle, Paragraph
    headers = tv["h"]; rows = tv["r"]; widths = tv.get("w")
    data = [[Paragraph(h, cell_h_s) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(v).replace("&", "&amp;"), cell_s) for v in row])
    col_w = None
    if widths:
        total = 17*cm
        col_w = [w*total for w in widths]
    t = Table(data, colWidths=col_w, repeatRows=1)
    style = [
        ('BACKGROUND', (0,0), (-1,0), C("1D4ED8")),
        ('GRID', (0,0), (-1,-1), 0.5, C("CBD5E1")),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
    ]
    for ri in range(1, len(data)):
        if ri % 2 == 0:
            style.append(('BACKGROUND', (0,ri), (-1,ri), C("F1F5F9")))
    t.setStyle(TableStyle(style))
    return t


if __name__ == "__main__":
    render_docx(DOCX_PATH)
    print("DOCX 生成完成：", DOCX_PATH)
    render_pdf(PDF_PATH)
    print("PDF  生成完成：", PDF_PATH)
